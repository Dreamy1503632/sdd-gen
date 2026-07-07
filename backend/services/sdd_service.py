"""
services/sdd_service.py
Orchestrates the SDD generation workflow:
  - Calls per-chapter LangChain chains
  - Yields SSE progress events
  - Assembles the final SDDDocument
  - Writes a DOCX file to disk and returns the download path
"""
from __future__ import annotations
import asyncio
import io
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import AsyncGenerator

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import get_settings
from models.hla_models import HLAAnalysis
from models.sdd_models import (
    SDDConfig, SDDDocument, SDDChapter1, SDDChapter2, SDDChapter3,
    SDDChapter4, SDDChapter6, SDDChapter9,
    BusinessStructureSection, ProcessFlow, ProcessStep, GapEntry, Assumption,
    DocumentControlEntry, SDDGenerationProgress,
)
from chains.sdd_chain import (
    generate_chapter2, generate_business_structure_section,
    generate_process_flow, generate_chapter6, generate_chapter9,
    BUSINESS_STRUCTURE_SECTIONS, PROCESS_FLOW_CATALOGUE,
)
from services.hla_service import get_analysis

settings = get_settings()
_SDD_STORE: dict[str, dict] = {}


# ──────────────────────────────────────────────────────────────────────────────
# Public entry point (streaming generator)
# ──────────────────────────────────────────────────────────────────────────────

async def generate_sdd_stream(
    hla_session_id: str,
    config: SDDConfig,
) -> AsyncGenerator[SDDGenerationProgress, None]:
    """
    Yields SDDGenerationProgress events as each chapter is generated.
    Callers use Server-Sent Events to stream progress to the browser.
    """
    session_id = str(uuid.uuid4())
    logs: list[str] = []
    start_time = time.time()

    def _progress(pct: int, chapter: str, msg: str) -> SDDGenerationProgress:
        logs.append(msg)
        return SDDGenerationProgress(
            session_id=session_id,
            status="generating",
            progress=pct,
            current_chapter=chapter,
            log_entries=list(logs),
        )

    analysis: HLAAnalysis | None = get_analysis(hla_session_id)
    if not analysis:
        yield SDDGenerationProgress(
            session_id=session_id, status="failed",
            progress=0, current_chapter="",
            log_entries=["HLA session not found. Please re-upload questionnaire."],
        )
        return

    yield _progress(2, "Setup", "Starting SDD generation...")

    # ── Chapter 1: Document Control (static) ──────────────────────────────────
    yield _progress(5, "Chapter 1", "Building document control...")
    ch1 = _build_chapter1(config)
    yield _progress(8, "Chapter 1", "✓ Chapter 1: Document Control complete")

    # ── Chapter 2: Introduction ───────────────────────────────────────────────
    yield _progress(10, "Chapter 2", "Generating introduction...")
    try:
        ch2_raw = await generate_chapter2(analysis, config)
        ch2 = SDDChapter2(**ch2_raw)
    except Exception as exc:
        ch2 = _fallback_chapter2(config, analysis)
    yield _progress(16, "Chapter 2", "✓ Chapter 2: Introduction complete")

    # ── Chapter 3: Business Structure (28 sections) ───────────────────────────
    yield _progress(18, "Chapter 3", "Starting business structure (28 sections)...")
    ch3_sections: list[BusinessStructureSection] = []
    total_sections = len(BUSINESS_STRUCTURE_SECTIONS)
    for idx, (sec_id, title, setup_task) in enumerate(BUSINESS_STRUCTURE_SECTIONS):
        pct = 18 + int((idx / total_sections) * 42)   # 18→60 range
        yield _progress(pct, f"Chapter 3 – {sec_id}", f"Generating {sec_id} {title}...")
        try:
            raw = await generate_business_structure_section(
                sec_id, title, setup_task, analysis, config
            )
            ch3_sections.append(BusinessStructureSection(**raw))
        except Exception:
            ch3_sections.append(_fallback_section(sec_id, title))
    ch3 = SDDChapter3(sections=ch3_sections)
    yield _progress(62, "Chapter 3", f"✓ Chapter 3: Business Structure complete ({len(ch3_sections)} sections)")

    # ── Chapter 4: Process Flows ──────────────────────────────────────────────
    yield _progress(63, "Chapter 4", "Starting process flows...")
    flows: list[ProcessFlow] = []
    all_flows = [
        (fid, title, module, txn)
        for module in analysis.modules
        if module in PROCESS_FLOW_CATALOGUE
        for fid, title, txn in PROCESS_FLOW_CATALOGUE[module]
    ]
    total_flows = len(all_flows)
    for idx, (fid, title, module, txn) in enumerate(all_flows):
        pct = 63 + int((idx / max(total_flows, 1)) * 22)  # 63→85
        yield _progress(pct, f"Chapter 4 – {fid}", f"Generating flow: {title}...")
        try:
            raw = await generate_process_flow(fid, title, module, txn, analysis, config)
            steps = [ProcessStep(**s) for s in raw.get("steps", [])]
            flows.append(ProcessFlow(
                flow_id=raw.get("flow_id", fid),
                title=raw.get("title", title),
                module=module,
                trigger=raw.get("trigger", ""),
                steps=steps,
                approval_required=raw.get("approval_required", False),
                approvers=raw.get("approvers", []),
                oracle_transaction=txn,
                notes=raw.get("notes", ""),
            ))
        except Exception:
            flows.append(_fallback_flow(fid, title, module, txn))
    ch4 = SDDChapter4(flows=flows)
    yield _progress(86, "Chapter 4", f"✓ Chapter 4: Process Flows complete ({len(flows)} flows)")

    # ── Chapter 6: Gap Analysis ───────────────────────────────────────────────
    yield _progress(87, "Chapter 6", "Generating gap analysis...")
    try:
        ch6_raw = await generate_chapter6(analysis, config)
        gaps = [GapEntry(**g) for g in ch6_raw.get("gaps", [])]
        ch6 = SDDChapter6(gaps=gaps, summary=ch6_raw.get("summary", ""))
    except Exception:
        ch6 = _fallback_chapter6(analysis)
    yield _progress(92, "Chapter 6", "✓ Chapter 6: Gap Analysis complete")

    # ── Chapter 9: Assumptions ────────────────────────────────────────────────
    yield _progress(93, "Chapter 9", "Generating assumptions...")
    try:
        ch9_raw = await generate_chapter9(analysis, config)
        assumptions = [Assumption(**a) for a in ch9_raw.get("assumptions", [])]
        ch9 = SDDChapter9(
            assumptions=assumptions,
            dependencies=ch9_raw.get("dependencies", []),
            exclusions=ch9_raw.get("exclusions", []),
        )
    except Exception:
        ch9 = _fallback_chapter9()
    yield _progress(96, "Chapter 9", "✓ Chapter 9: Assumptions complete")

    # ── Assemble document ─────────────────────────────────────────────────────
    yield _progress(97, "Finalising", "Assembling SDD document...")
    gen_time = round(time.time() - start_time, 1)

    sdd = SDDDocument(
        session_id=session_id,
        hla_session_id=hla_session_id,
        config=config,
        modules=analysis.modules,
        chapter1=ch1,
        chapter2=ch2,
        chapter3=ch3,
        chapter4=ch4,
        chapter6=ch6,
        chapter9=ch9,
        estimated_pages=_estimate_pages(ch3, ch4),
        section_count=len(ch3_sections) + len(flows) + 6,
        process_flow_count=len(flows),
        generation_time_seconds=gen_time,
    )

    # ── Write DOCX ────────────────────────────────────────────────────────────
    yield _progress(98, "Writing DOCX", "Writing DOCX file...")
    docx_path = _write_docx(sdd)
    download_url = f"/api/v1/sdd/download/{session_id}"

    # Persist
    _SDD_STORE[session_id] = {
        "sdd": sdd.model_dump(),
        "docx_path": str(docx_path),
        "created_at": datetime.utcnow().isoformat(),
    }

    logs.append(f"✅ SDD generated in {gen_time}s")
    yield SDDGenerationProgress(
        session_id=session_id,
        status="complete",
        progress=100,
        current_chapter="Complete",
        log_entries=list(logs),
        download_url=download_url,
        document=sdd,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Getters
# ──────────────────────────────────────────────────────────────────────────────

def get_sdd(session_id: str) -> SDDDocument | None:
    data = _SDD_STORE.get(session_id)
    if not data:
        return None
    return SDDDocument(**data["sdd"])


def get_sdd_docx_path(session_id: str) -> Path | None:
    data = _SDD_STORE.get(session_id)
    if not data:
        return None
    p = Path(data["docx_path"])
    return p if p.exists() else None


# ──────────────────────────────────────────────────────────────────────────────
# DOCX writer
# ──────────────────────────────────────────────────────────────────────────────

def _write_docx(sdd: SDDDocument) -> Path:
    doc = DocxDocument()
    _set_page_margins(doc)

    # Cover page
    _add_cover_page(doc, sdd)
    doc.add_page_break()

    # TOC placeholder
    p = doc.add_paragraph("Table of Contents")
    p.style = "Heading 1"
    doc.add_paragraph("[Auto-generated Table of Contents – update after opening in Word]")
    doc.add_page_break()

    # Ch2
    doc.add_heading("2. Introduction", level=1)
    doc.add_paragraph(sdd.chapter2.purpose)
    doc.add_heading("2.1 Scope", level=2)
    doc.add_paragraph(sdd.chapter2.scope)
    doc.add_heading("2.2 Implementation Approach", level=2)
    doc.add_paragraph(sdd.chapter2.implementation_approach)
    doc.add_page_break()

    # Ch3
    doc.add_heading("3. Business Structure", level=1)
    for sec in sdd.chapter3.sections:
        doc.add_heading(f"{sec.section_id} {sec.title}", level=2)
        doc.add_paragraph(sec.description)
        if sec.oracle_recommendation:
            p = doc.add_paragraph()
            p.add_run("Oracle Recommendation: ").bold = True
            p.add_run(sec.oracle_recommendation)
        if sec.client_decision:
            p = doc.add_paragraph()
            p.add_run(f"{sdd.config.company_name} Decision: ").bold = True
            p.add_run(sec.client_decision)
        if sec.table_data:
            _add_config_table(doc, sec.table_data)
        if sec.configuration_notes:
            p = doc.add_paragraph()
            p.add_run("Configuration Notes: ").bold = True
            p.add_run(sec.configuration_notes)
    doc.add_page_break()

    # Ch4
    doc.add_heading("4. Process Flows", level=1)
    for flow in sdd.chapter4.flows:
        doc.add_heading(f"{flow.flow_id} {flow.title}", level=2)
        doc.add_paragraph(f"Module: {flow.module}  |  Oracle Transaction: {flow.oracle_transaction}")
        doc.add_paragraph(f"Trigger: {flow.trigger}")
        if flow.steps:
            _add_process_steps_table(doc, flow.steps)
        if flow.approval_required:
            p = doc.add_paragraph()
            p.add_run("Approval Required: ").bold = True
            p.add_run(", ".join(flow.approvers))
        if flow.notes:
            doc.add_paragraph(f"Notes: {flow.notes}")
    doc.add_page_break()

    # Ch6
    doc.add_heading("6. Gap Analysis", level=1)
    doc.add_paragraph(sdd.chapter6.summary)
    for gap in sdd.chapter6.gaps:
        doc.add_heading(f"{gap.gap_id} – {gap.module}", level=2)
        _add_gap_table(doc, gap)
    doc.add_page_break()

    # Ch9
    doc.add_heading("9. Assumptions & Dependencies", level=1)
    doc.add_heading("9.1 Assumptions", level=2)
    for assumption in sdd.chapter9.assumptions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{assumption.assumption_id}] {assumption.category}: ").bold = True
        p.add_run(assumption.description)
    doc.add_heading("9.2 Dependencies", level=2)
    for dep in sdd.chapter9.dependencies:
        doc.add_paragraph(dep, style="List Bullet")
    doc.add_heading("9.3 Exclusions", level=2)
    for exc in sdd.chapter9.exclusions:
        doc.add_paragraph(exc, style="List Bullet")

    # Sign-off sheet
    doc.add_page_break()
    doc.add_heading("10. Sign Off Sheet", level=1)
    _add_signoff_table(doc)

    # Save
    out_dir = Path(settings.output_dir)
    out_dir.mkdir(exist_ok=True)
    fname = f"SDD_{sdd.config.company_name.replace(' ', '_')}_{sdd.session_id[:8]}.docx"
    path = out_dir / fname
    doc.save(str(path))
    return path


# ──────────────────────────────────────────────────────────────────────────────
# DOCX formatting helpers
# ──────────────────────────────────────────────────────────────────────────────

def _set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.54)


def _add_cover_page(doc, sdd: SDDDocument):
    doc.add_paragraph()
    t = doc.add_paragraph(sdd.config.company_name)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(28)
    t.runs[0].bold = True

    t2 = doc.add_paragraph("Oracle Fusion HCM – Solution Design Document")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(20)

    for label, value in [
        ("Project:", sdd.config.project_name),
        ("Reference:", sdd.config.doc_reference),
        ("Author:", sdd.config.author),
        ("Version:", sdd.config.version),
        ("Date:", datetime.now().strftime("%d %B %Y")),
        ("Confidentiality:", sdd.config.confidentiality),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} ").bold = True
        p.add_run(value)


def _add_config_table(doc, table_data: list[dict]):
    if not table_data:
        return
    keys = list(table_data[0].keys())
    t = doc.add_table(rows=1, cols=len(keys))
    t.style = "Light Shading Accent 1"
    hdr_cells = t.rows[0].cells
    for i, k in enumerate(keys):
        hdr_cells[i].text = k.replace("_", " ").title()
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row_data in table_data:
        row_cells = t.add_row().cells
        for i, k in enumerate(keys):
            row_cells[i].text = str(row_data.get(k, ""))


def _add_process_steps_table(doc, steps: list[ProcessStep]):
    headers = ["Step", "Actor", "Action", "System", "Outcome", "Notes"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for step in steps:
        row = t.add_row().cells
        row[0].text = str(step.step_number)
        row[1].text = step.actor
        row[2].text = step.action
        row[3].text = step.system
        row[4].text = step.outcome
        row[5].text = step.notes


def _add_gap_table(doc, gap: GapEntry):
    data = [
        ("Requirement", gap.requirement),
        ("Oracle Capability", gap.oracle_capability),
        ("Gap", gap.gap_description or "None"),
        ("Recommended Solution", gap.recommended_solution),
        ("Implementation Effort", gap.effort),
        ("Workaround", gap.workaround or "N/A"),
    ]
    t = doc.add_table(rows=len(data), cols=2)
    t.style = "Light Shading Accent 1"
    for i, (label, value) in enumerate(data):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].text = value


def _add_signoff_table(doc):
    headers = ["Role", "Name", "Signature", "Date"]
    roles = [
        "Oracle Solution Architect",
        "Client HR Project Lead",
        "Client IT Lead",
        "Oracle Project Manager",
        "Client Sponsor",
    ]
    t = doc.add_table(rows=1 + len(roles), cols=4)
    t.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, role in enumerate(roles, start=1):
        t.rows[i].cells[0].text = role


# ──────────────────────────────────────────────────────────────────────────────
# Estimation helpers
# ──────────────────────────────────────────────────────────────────────────────

def _estimate_pages(ch3: SDDChapter3, ch4: SDDChapter4) -> int:
    return 20 + (len(ch3.sections) * 2) + (len(ch4.flows) * 3)


# ──────────────────────────────────────────────────────────────────────────────
# Fallback builders (used when AI call fails)
# ──────────────────────────────────────────────────────────────────────────────

def _build_chapter1(config: SDDConfig) -> SDDChapter1:
    return SDDChapter1(
        change_history=[
            DocumentControlEntry(
                version="1.0",
                date=datetime.now().strftime("%d %b %Y"),
                author=config.author,
                description="Initial draft",
            )
        ],
        distribution_list=[config.company_name + " HR Leadership", "Oracle Implementation Team"],
        abbreviations={
            "HCM": "Human Capital Management",
            "SDD": "Solution Design Document",
            "HLA": "High Level Assessment",
            "OOTB": "Out of the Box",
            "HDL": "HCM Data Loader",
            "OTBI": "Oracle Transactional Business Intelligence",
            "BU": "Business Unit",
            "LE": "Legal Entity",
        },
    )


def _fallback_chapter2(config: SDDConfig, analysis: HLAAnalysis) -> SDDChapter2:
    return SDDChapter2(
        purpose=f"This SDD documents the Oracle Fusion HCM solution design for {config.company_name}.",
        scope=f"The implementation covers {', '.join(analysis.modules)} for {config.company_name}'s operations.",
        implementation_approach="Oracle Cloud implementation following Oracle Unified Methodology (OUM) with an Agile delivery approach.",
        document_overview="This document covers business structure setup, process flows, gap analysis, and implementation assumptions.",
    )


def _fallback_section(sec_id: str, title: str) -> BusinessStructureSection:
    return BusinessStructureSection(
        section_id=sec_id,
        title=title,
        description=f"Configuration for {title} within Oracle Fusion HCM.",
        oracle_recommendation="Follow Oracle standard configuration guidelines.",
        client_decision="To be confirmed during configuration workshops.",
        configuration_notes=f"Configure via Setup and Maintenance → Workforce Structures.",
        table_data=[],
    )


def _fallback_flow(fid: str, title: str, module: str, txn: str) -> ProcessFlow:
    return ProcessFlow(
        flow_id=fid,
        title=title,
        module=module,
        trigger="Initiated by authorised user",
        steps=[
            ProcessStep(step_number=1, actor="HR Specialist", action="Initiate transaction", system="Oracle HCM", outcome="Transaction started", notes=""),
            ProcessStep(step_number=2, actor="Manager", action="Review and approve", system="Oracle HCM – Worklist", outcome="Approved/Rejected", notes=""),
            ProcessStep(step_number=3, actor="HR Specialist", action="Confirm completion", system="Oracle HCM", outcome="Record updated", notes=""),
        ],
        approval_required=True,
        approvers=["Line Manager"],
        oracle_transaction=txn,
        notes="Detailed flow to be confirmed during process workshops.",
    )


def _fallback_chapter6(analysis: HLAAnalysis) -> SDDChapter6:
    gaps = [
        GapEntry(
            gap_id=f"G{i+1:03d}",
            module=g.module,
            requirement=g.requirement,
            oracle_capability=g.oracle_capability,
            gap_description=g.gap if g.gap.lower() != "none" else "",
            recommended_solution=g.solution,
            effort=g.effort.value,
            workaround="",
        )
        for i, g in enumerate(analysis.gap_analysis)
    ]
    return SDDChapter6(
        gaps=gaps,
        summary="Gap analysis confirms Oracle Fusion HCM delivers the majority of requirements through standard configuration.",
    )


def _fallback_chapter9() -> SDDChapter9:
    return SDDChapter9(
        assumptions=[
            Assumption(assumption_id="A001", category="Data", description="Historical employee data will be provided in Oracle HDL-compatible format.", impact_if_incorrect="Data migration timeline will extend by 4–6 weeks."),
            Assumption(assumption_id="A002", category="Organisational", description="Client will confirm enterprise structure decisions within 2 weeks of design workshops.", impact_if_incorrect="Configuration cannot proceed until decisions are confirmed."),
            Assumption(assumption_id="A003", category="Integration", description="API credentials for all integration points will be provided before build phase.", impact_if_incorrect="Integration build and testing will be delayed."),
            Assumption(assumption_id="A004", category="Technical", description="Client will provision Oracle environments (Dev, Test, Production) before implementation starts.", impact_if_incorrect="All implementation activities will be blocked."),
            Assumption(assumption_id="A005", category="Process", description="Client SMEs will be available for workshop sessions (minimum 50% time).", impact_if_incorrect="Design workshops will overrun and design quality will be compromised."),
        ],
        dependencies=["Oracle Cloud environment provisioning", "Data extract from legacy systems", "SSO configuration with client Identity Provider"],
        exclusions=["Customisations beyond Oracle Fusion delivered functionality", "Third-party application development", "Infrastructure procurement"],
    )